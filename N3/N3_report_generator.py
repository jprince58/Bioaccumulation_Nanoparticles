#!/usr/bin/env python3
# -*- coding: utf-8 -*-

vn_report_generator=1.0

import os
import numpy as np
import matplotlib.pyplot as plt
import docx
from docx.shared import Pt
from datetime import datetime
import matplotlib.animation as anim
from matplotlib.animation import FuncAnimation

"Commenting out correct plot generator function while lienar fitting functionality turned off"
def plot_generator(c_set,parameter_combos_count,parameter_matrix,new_count_number,vn_N2,vn_Main_Code,vn_parameter_matrix_generator,vn_parameter_checker,vn_csv_generator,vn_method_of_lines,vn_RJ,perc_acc_matrix,vn_linear_fitting,machine_number,internal_export_path):
# def plot_generator(c_set,parameter_combos_count,parameter_matrix,new_count_number,vn_N2,vn_Main_Code,vn_parameter_matrix_generator,vn_parameter_checker,vn_csv_generator,vn_method_of_lines,vn_RJ,machine_number,internal_export_path):
    """Static Plotting (Exported to Word Document)"""
    report=docx.Document()
    report.add_heading(f'Results from N2 Run #{new_count_number}-{machine_number}',0)
    now = datetime.now()
    dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
    date_time_line=report.add_paragraph('Date and Time Report Generated:  ')
    date_time_line.add_run(dt_string)
    report.add_paragraph(f'Machine number ran on- {machine_number}')
    report.add_paragraph(f'N2 version: {vn_N2}')
    report.add_paragraph(f'Main Code version: {vn_Main_Code}')
    report.add_paragraph(f'Parameter Matrix Generator version: {vn_parameter_matrix_generator}')
    report.add_paragraph(f'Parameter Checker version: {vn_parameter_checker}')
    report.add_paragraph(f'CSV Generator version: {vn_csv_generator}')
    report.add_paragraph(f'Method of Lines version: {vn_method_of_lines}')
    report.add_paragraph(f'Residual-Jacobian Calculator version: {vn_RJ}')
    report.add_paragraph(f'Report Generator version: {vn_report_generator}')
    #report.add_paragraph(f'Linear Approximator version: {vn_linear_fitting}')
    style=report.styles['Normal']
    font=style.font
    font.name='Arial'
    font.size=Pt(9)  
        
     
    for pc_i in np.arange(0,parameter_combos_count,1): #Begin for loop to plot the different model paramters using MOL 
        cb=c_set[pc_i][0] #Grab current bound concentration data to plot
        cu=c_set[pc_i][1] #Grab current unbound concentration data to plot
        pot=c_set[pc_i][2] #Grab current potential concentration data to plot
        average_uconc_overtime=c_set[pc_i][3] #Grab current average unbound concentration data to plot
        average_bconc_overtime=c_set[pc_i][4] #Grab current average bound concentration data to plot
        average_tconc_overtime=c_set[pc_i][5] #Grab current average total concentration data to plot
        lognorm_tconc_overtime=c_set[pc_i][6] #Grab current change in log&normalized total concentration data to plot
        t=c_set[pc_i][7] #Grab time-vector for this parameter set for plotting
        nt=c_set[pc_i][8] #Grab number of time points for this parameter set for plotting
        x=c_set[pc_i][9] #Grab the position-vector for this parameter set for plotting
        average_tconc_ss=c_set[pc_i][13]
        break_paragraph=report.add_paragraph('___________')
        break_paragraph.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
        report.add_heading('Parameter Set %i'%pc_i,1)
        para1=report.add_paragraph(f'Step-size (h) : {parameter_matrix[pc_i,0]}     ')
        para1.add_run(f'Initial time (t1) : {parameter_matrix[pc_i,2]}     ')
        para1.add_run(f'Final time (t2) : {parameter_matrix[pc_i,3]}     ')
        para1.add_run(f'Mesh size (nx) : {parameter_matrix[pc_i,4]}')
        para2=report.add_paragraph(f'Dimensionless ratio of diffusivity (gamma) : {parameter_matrix[pc_i,5]}     ')
        para2.add_run(f'Dimensionless ratio of potential (beta): {parameter_matrix[pc_i,12]}')
        para3=report.add_paragraph(f'Dimensionless forward rate constant (F): {parameter_matrix[pc_i,6]}     ')
        para3.add_run(f'Nanoparticle Binding Constant (K): {parameter_matrix[pc_i,7]}')
        para4=report.add_paragraph(f'Ratio of binding sites to supernatant concentration (epsilon): {parameter_matrix[pc_i,8]}     ')
        para4.add_run(f'Ratio of Nanoparticle to Biofilm charge (upsilon): {parameter_matrix[pc_i,10]}')
        para5=report.add_paragraph(f'Partition Coeffecient of Nanoparticle from Water to Biofilm (Kp): {parameter_matrix[pc_i,11]}     ')
        para5.add_run(f'NP size contribution to charge (omega): {parameter_matrix[pc_i,9]}')
        para6=report.add_paragraph(f'Tolerance: {parameter_matrix[pc_i,1]}')
        para6.add_run(f'Average Total Dimensionless NP Concentration @ SS: {average_tconc_ss}')
        
        # %%Find relelvant maximums and minimums
        upper_1 = np.amax(cu)*1.1 #Upper bound on Unbound Concentration
        upper_2 = np.amax(average_uconc_overtime)*1.1 #Upper Bound on Average Unbound Concentration
        upper_3 = np.amax(average_bconc_overtime)*1.1 #Upper Bound on Average Bound Concentration
        upper_4 = np.amax(cb)*1.1 #Upper bound on Bound Concentration
        upper_5 = np.amax(average_tconc_overtime)*1.1 #Upper Bound on Average Total Concentration
        upper_6 = np.amax(lognorm_tconc_overtime)*1.1 #Upper Bound on Average log&normalized total concentration
        lower_6 = np.amin(lognorm_tconc_overtime)*0.9 #Lower Bound on log&normalized total concentration
        upper_9 = np.amax(pot)*1.1 #Upper bound on potential
        lower_9 = np.amin(pot)*1.1 #Lower Bound on potential
        
        # %%Unbound
        #tindex_u=np.array([0,5,10,25,50,75,100,125,150,200,250]) for masnual control over timepoints plotted
        tp_u=10 #number of time points to plot

        """
        #Linear discretization of plotted timepionts
        space_u=int((nt-1)/tp_u) #Linear discreitzation of timepoints
        tindex_u=np.arange(0,nt,space_u) #Linear discreitization of timepoints
        for i_u in tindex_u:
            cc_u=cu[:,i_u]
            ti_u=round(t[i_u],5)
            plt.plot(x,cc_u,label='t={}'.format(ti_u))
        """
        #Logarthmic discreitzation of plotted timepoints
        lognt_u=np.log10(nt) #Logarthmic timepoints
        logspace_u=round((lognt_u)/tp_u,10) #Logarthmic timepoints
        logtindex_u=np.arange(0,lognt_u,logspace_u) #Logarthmic timepoints
        plt.figure(7*pc_i+0)
        for logi_u in logtindex_u:
            i_u=int(10**logi_u)
            cc_u=cu[:,i_u]
            ti_u=round(t[i_u],5)
            plt.plot(x,cc_u,label='t={}'.format(ti_u))
        
        plt.xlim(left=0,right=1)
        plt.ylim(bottom=0,top=upper_1)
        plt.xlabel('Position',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Dimensionless Unbound Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        plt.legend(loc=(0.1,0.1))
        unbound_filename_partial=f'Unboundplot{pc_i}.png'
        unbound_filename_full=os.path.join(internal_export_path,unbound_filename_partial)
        plt.savefig(unbound_filename_full)
        plt.close()
       
        
        # %%Bound
        tp_b=10 #number of time points to plot

    
        # #Linear discretization of plotted timepionts
        # space_b=int((nt-1)/tp_b) #Linear discreitzation of timepoints
        # tindex_b=np.arange(0,nt,space_b) #Linear discreitzation of timepoints
        # for i_b in tindex_b:
        #     cc_b=cb[:,i_b]
        #     ti_b=round(t[i_b],5)
        #     plt.plot(x,cc_b,label='t={}'.format(ti_b))

        #Logarthmic discreitzation of plotted timepoints
        lognt_b=np.log10(nt) #Logarthmic timepoints
        logspace_b=round((lognt_b)/tp_b,10) #Logarthmic timepoints
        logtindex_b=np.arange(0,lognt_b,logspace_b) #Logarthmic timepoints
        plt.figure(7*pc_i+1)
        for logi_b in logtindex_b:
            i_b=int(10**logi_b)
            cc_b=cb[:,i_b]
            ti_b=round(t[i_b],5)
            plt.plot(x,cc_b,label='t={}'.format(ti_b))
   
        plt.xlim(left=0,right=1)
        plt.ylim(bottom=0,top=upper_4)
        plt.xlabel('Position',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Dimensionless Bound Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        plt.legend(loc=(0.1,0.1))
        bound_filename_partial=f'Boundplot{pc_i}.png'
        bound_filename_full=os.path.join(internal_export_path,bound_filename_partial)
        plt.savefig(bound_filename_full)
        plt.close()
        pics_paragraph1=report.add_paragraph()
        pic1=pics_paragraph1.add_run()
        pic1.add_picture(unbound_filename_full, width=docx.shared.Inches(3))
        pic2=pics_paragraph1.add_run()
        pic2.add_picture(bound_filename_full, width=docx.shared.Inches(3))
        
        # %%Potential
        tp_p=10 #number of time points to plot
        
        #Logarthmic discreitzation of plotted timepoints
        lognt_p=np.log10(nt) #Logarthmic timepoints
        logspace_p=round((lognt_p)/tp_p,10) #Logarthmic timepoints
        logtindex_p=np.arange(0,lognt_p,logspace_p) #Logarthmic timepoints
        plt.figure(7*pc_i+2)
        for logi_p in logtindex_p:
            i_p=int(10**logi_p)
            cc_p=pot[:,i_p]
            ti_p=round(t[i_p],5)
            plt.plot(x,cc_p,label='t={}'.format(ti_p))
        
        plt.xlim(left=0,right=1)
        plt.ylim(bottom=lower_9,top=upper_9)
        plt.xlabel('Position',fontsize=14)
        plt.ylabel('Dimensionless Potential',fontsize=14)
        plt.title('Dimensionless Potential plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        plt.legend(loc=(0.1,0.1))
        potential_filename_partial=f'Potentialplot{pc_i}.png'
        potential_filename_full=os.path.join(internal_export_path,potential_filename_partial)
        plt.savefig(potential_filename_full)
        plt.close()
        
        # %%Unbound NP Average Concetration Overtime
        plt.figure(7*pc_i+3)
        plt.plot(t,average_uconc_overtime)
        plt.xlim(left=parameter_matrix[pc_i,2],right=parameter_matrix[pc_i,3])
        #plt.xlim(left=0,right=0.0005) #Manual Override of automatic x-axis limits
        plt.ylim(bottom=0,top=upper_2)
        plt.xlabel('Time',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Average Dimensionless Unbound Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        avgunbound_filename_partial=f'AvgUnboundplot{pc_i}.png'
        avgunbound_filename_full=os.path.join(internal_export_path,avgunbound_filename_partial)
        plt.savefig(avgunbound_filename_full)
        plt.close()
        pics_paragraph2=report.add_paragraph()
        pic3=pics_paragraph2.add_run()
        pic3.add_picture(potential_filename_full, width=docx.shared.Inches(3))
        pic4=pics_paragraph2.add_run()
        pic4.add_picture(avgunbound_filename_full, width=docx.shared.Inches(3))
        
        #%%Bound NP Average Concetration Overtime
        plt.figure(7*pc_i+4)
        plt.plot(t,average_bconc_overtime)
        plt.xlim(left=parameter_matrix[pc_i,2],right=parameter_matrix[pc_i,3])
        #plt.xlim(left=0,right=0.0005) #Manual Override of automatic x-axis limits
        plt.ylim(bottom=0,top=upper_3)
        plt.xlabel('Time',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Average Dimensionless Bound Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        avgbound_filename_partial=f'AvgBoundplot{pc_i}.png'
        avgbound_filename_full=os.path.join(internal_export_path,avgbound_filename_partial)
        plt.savefig(avgbound_filename_full)
        plt.close()
        
        #%%Total NP Average Concetration Overtime
        plt.figure(7*pc_i+5)
        plt.plot(t,average_tconc_overtime)
        plt.xlim(left=parameter_matrix[pc_i,2],right=parameter_matrix[pc_i,3])
        #plt.xlim(left=0,right=0.0005) #Manual Override of automatic x-axis limits
        plt.ylim(bottom=0,top=upper_5)
        plt.xlabel('Time',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Average Dimensionless Total Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        avgtotal_filename_partial=f'AvgTotalplot{pc_i}.png'
        avgtotal_filename_full=os.path.join(internal_export_path,avgtotal_filename_partial)
        plt.savefig(avgtotal_filename_full)
        plt.close()
        
        pics_paragraph3=report.add_paragraph()
        pic5=pics_paragraph3.add_run()
        pic5.add_picture(avgbound_filename_full, width=docx.shared.Inches(3))
        pic6=pics_paragraph3.add_run()
        pic6.add_picture(avgtotal_filename_full, width=docx.shared.Inches(3))
        
        # # %% Model fit Plot
        # pics_paragraph4=report.add_paragraph()
        # modelfit_filename_partial=f'Modelfitplot{pc_i}.png'
        # modelfit_filename_full=os.path.join(internal_export_path,modelfit_filename_partial)
        # pic4=pics_paragraph4.add_run()
        # pic4.add_picture(modelfit_filename_full, width=docx.shared.Inches(6))
        
        # %%Log&Normalized NP Average Concetration Overtime
        plt.figure(7*pc_i+6)
        plt.plot(t,lognorm_tconc_overtime)
        plt.xlim(left=parameter_matrix[pc_i,2],right=parameter_matrix[pc_i,3])
        #plt.xlim(left=0,right=0.0005) #Manual Override of automatic x-axis limits
        plt.ylim(bottom=lower_6,top=upper_6)
        plt.xlabel('Time',fontsize=14)
        plt.ylabel('Log of Normalized Concentration',fontsize=14)
        plt.title('First-order Plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        lognorm_filename_partial=f'Firstorder{pc_i}.png'
        lognorm_filename_full=os.path.join(internal_export_path,lognorm_filename_partial)
        plt.savefig(lognorm_filename_full)
        plt.close()
 
        # # %%Unbound Concentration Animation
        # unbound_anim_fig=plt.figure()
        # unbound_anim_plot=plt.plot([])
        # unbound_anim_holder=unbound_anim_plot[0]
        # plt.xlim(left=0,right=1)
        # plt.ylim(bottom=0,top=upper_1)
        # plt.xlabel('Position',fontsize=14)
        # plt.ylabel('Dimensionless Concentration',fontsize=14)
        # plt.title('Dimensionless Unbound Concentration',fontsize=16)
        # plt.xticks(fontsize=12)
        # plt.yticks(fontsize=12)
        # tp_u_anim=100 #number of time points to plot for the animation for unbound concntration
        # space_u_anim=int((nt-1)/tp_u_anim) #space between timepoints    
        # def unbound_animate(frame):
        #     #update plot
        #     c_u_plot=cu[:,frame*space_u_anim]
        #     unbound_anim_holder.set_data((x,c_u_plot))
        # unbound_anim=anim.FuncAnimation(unbound_anim_fig,unbound_animate,frames=tp_u_anim,interval=100)
        # unbound_anim_filename_partial=f'unboun_anim{pc_i}.gif'
        # unbound_anim_filename_full=os.path.join(internal_export_path,unbound_anim_filename_partial)
        # unbound_anim.save(unbound_anim_filename_full)
        # #Only need these lines if log plot is turned off

        
        # %%Total NP Concentration vs Time Animation
        totCvt_anim_fig=plt.figure()
        totCvt_anim_plot=plt.plot([])
        totCvt_anim_holder=totCvt_anim_plot[0]
        plt.xlim(left=parameter_matrix[pc_i,2],right=parameter_matrix[pc_i,3])
        plt.ylim(bottom=0,top=upper_5)
        plt.xlabel('Time',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Average Dimensionless Total Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        tp_totCvt_anim=100 #number of time points to plot for the animation for unbound concntration
        space_totCvt_anim=int((nt-1)/tp_totCvt_anim) #space between timepoints    
        def totCvt_animate(frame):
            #update plot
            totCvt_anim_holder.set_data((t[0:frame*space_totCvt_anim],average_tconc_overtime[0:frame*space_totCvt_anim]))
        totCvt_anim=anim.FuncAnimation(totCvt_anim_fig,totCvt_animate,frames=tp_totCvt_anim,interval=100)
        totCvt_anim_filename_partial=f'totCvt_anim{pc_i}.gif'
        totCvt_anim_filename_full=os.path.join(internal_export_path,totCvt_anim_filename_partial)
        # totCvt_anim.save(totCvt_anim_filename_full)

        pics_paragraph4=report.add_paragraph()
        pic7=pics_paragraph4.add_run()
        pic7.add_picture(lognorm_filename_full, width=docx.shared.Inches(3))
        pic8=pics_paragraph4.add_run()
        pic8.add_picture(totCvt_anim_filename_full, width=docx.shared.Inches(3))      
        plt.close()
        
        # %% Plot Approximation for Total NP conc Overtime
        linear_filename_partial=f'Linearplot{pc_i}.png'
        linear_filename_full=os.path.join(internal_export_path,linear_filename_partial)
        #pics_paragraph6=report.add_paragraph() commented out  when log plots out
        #pic9=pics_paragraph6.add_run() commented out  when log plots out
        pics_paragraph5=report.add_paragraph()
        pic9=pics_paragraph5.add_run()
        pic9.add_picture(linear_filename_full,width=docx.shared.Inches(3))
        log_filename_partial=f'Logplot{pc_i}.png'
        log_filename_full=os.path.join(internal_export_path,log_filename_partial)
        pic10=pics_paragraph5.add_run() #Added when log plot off
        pic10.add_picture(log_filename_full,width=docx.shared.Inches(3))
        
        # # %% Add Table for Fit
        # perc_acc_table=perc_acc_matrix[pc_i][0]
        # [table_rows,table_columns]=perc_acc_table.shape
        # table1=report.add_table(rows=table_rows+1, cols=table_columns)
        # row=table1.rows[0]
        # row.cells[0].text='Percent Accumulated'
        # row.cells[1].text='Time for Model'
        # row.cells[2].text='Time for Approximation'
        # row.cells[3].text='Percent Error'
        # i_v = np.arange(0,table_rows,1) #index for rows of percent accumulation table
        # j_v = np.arange(0,table_columns,1) #inde for columns of percent accumulation table
        # for i in i_v:
        #     for j in j_v:
        #         cell=table1.cell(i+1,j)
        #         cell.text=str(perc_acc_table[i,j])
            
    # sherwood_filename_partial=f'Sherwdoodplot{pc_i}.png'
    # sherwood_filename_full=os.path.join(internal_export_path,sherwood_filename_partial)
    # pics_paragraph6=report.add_paragraph()
    # pic11=pics_paragraph6.add_run()
    # pic11.add_picture(sherwood_filename_full,width=docx.shared.Inches(3))
    # report.add_paragraph(f'Slope/interecpt/R^2 is: {perc_acc_matrix[0][4]}')
    # #para7=report.add_paragraph(f'R^2={perc_acc_matrix[0][3]}     ')
    
    # logsherwood_filename_partial=f'LogSherwdoodplot{pc_i}.png'
    # logsherwood_filename_full=os.path.join(internal_export_path,logsherwood_filename_partial)
    # pics_paragraph6=report.add_paragraph()
    # pic11=pics_paragraph6.add_run()
    # pic11.add_picture(logsherwood_filename_full,width=docx.shared.Inches(3))
    # report.add_paragraph(f'Slope/interecpt/R^2 is: {perc_acc_matrix[0][5]}')
    
    # %% Add Table for Multiple Linear Regression
    # coef_multreg=perc_acc_matrix[0][1]
    # intercept_multreg=perc_acc_matrix[0][2]
    # table_rows=len(coef_multreg)+2
    # table_columns=2
    # table2=report.add_table(rows=table_rows, cols=table_columns)
    # row=table2.rows[0]
    # row.cells[0].text='Beta Parameter'
    # row.cells[1].text='Value'
    # row=table2.rows[1]
    # row.cells[0].text='Beta0'
    # row.cells[1].text=f'{intercept_multreg}'
    # i_v = np.arange(2,table_rows,1) #index for rows of percent accumulation table
    # for i in i_v:
    #     cell=table2.cell(i,1)
    #     cell.text=str(coef_multreg[i-2])
    #     cell=table2.cell(i,0)
    #     cell.text=f'Beta{i-1}'

        
        # %%
    return report

# %%
"""
Created on Sat Oct 31 18:13:53 2020

@author: joshuaprince

Purpose: Script to auto-generate report from data, including plotting of key figures and generation of an animated plot

Version 1.0

Changes from Version 0.1 to 1.0 (2/20/2022 1:30 am)
    Got initial ocde to run. Had to comment out most functionalities
    
"""