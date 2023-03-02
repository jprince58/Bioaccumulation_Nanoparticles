#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Jun 16 16:13:07 2020

@author: joshuaprince
"""


#Inputs Code Block
"""
This cell take in the inputs to this code. This does so by:
1a) takes in the user-specified dimensional values for the code, or 
1b) takes in user-specified dimensionless numbers then
2) defines the relevant inputs and outputs which go to the heart of the code 

This is the primary code one should be modifying when testing different physical systems for the model
"""

import time
import numpy as np
import scipy as sp
import matplotlib.pyplot as plt
from numpy import linalg
from numpy.linalg import solve
from scipy import sparse
from scipy.sparse import linalg
from scipy.sparse import csc_matrix
from matplotlib.backends.backend_pdf import PdfPages
import docx
from docx.shared import Pt
from matplotlib.animation import FuncAnimation

#Editing Code

#Start Timer
t_start=time.time()

def RJ(x,y,p):
    nx=len(x)-1 #Grab the mesh size for position
    ny=len(y)-2 #Grab number of y-points
    dx=1/nx #Calculate the distance between nodes (assumes domain is from 0 to 1)
    gam=p[0] #Grab the dimenionless ratio of diffusivities
    beta=p[1] #Grab the dimensionless ratio of potentials
    F=p[2] #Grab the dimensionless forward reaction rate constant
    Re=p[3] #Grab the dimensionless reverse reaction rate constant
    n=p[4] #Grab the hill coeffecient
    R=np.zeros(ny+2) #Initialize R
    J=np.zeros((ny+2,ny+2)) #Initialize J
    index=np.arange(0,ny+2) #Creater index vector
    for i in index:
        if i==0 :
            R[i]=y[i]-1
            J[i,i]=1;
        elif i%2==1 :
            R[i]=F*y[i-1]**n*(1-y[i])-Re*y[i]
            J[i,i]=-(F*y[i-1]**n+Re)
            J[i,i-1]=n*F*y[i-1]**(n-1)*(1-y[i])
        elif i==ny:
            l=int(i/2)
            R[i]=2*(y[i-2]-y[i])/dx**2*(1-x[l]**2+gam)-y[i]*(F*y[i]**(n-1)*(1-y[i+1])-2*beta)+Re*y[i+1]
            J[i,i]=(-2/dx**2)*(1-x[l]**2+gam)-(n*F*y[i]**(n-1)*(1-y[i-1]-2*beta))
            J[i,i-2]=2/dx**2*(1-x[l]**2+gam)
            J[i,i+1]=F*y[i]**n+Re
        elif i%2==0 and i!=0:
            l=int(i/2)
            R[i]=(y[i+2]-2*y[i]+y[i-2])/dx**2*(1-x[l]**2+gam)-(y[i+2]-y[i-2])/2/dx*(2*x[l]*(1-beta))-y[i]*(F*y[i]**(n-1)*(1-y[i+1])-2*beta)+Re*y[i+1]
            J[i,i]=(-2/dx**2)*(1-x[l]**2+gam)-(n*F*y[i]**(n-1)*(1-y[i-1]-2*beta))
            J[i,i+2]=1/dx**2*(1-x[l]**2+gam)-1/2/dx*(2*x[l]*(1-beta))
            J[i,i-2]=1/dx**2*(1-x[l]**2+gam)+1/2/dx*(2*x[l]*(1-beta))
            J[i,i+1]=F*y[i]**n+Re
        else:
            print('Uh oh')
    return (R,J)



def MOL(t,x,y,h,p,tol):
    yw=y[:,0] #Initalize the working concentration vector
    index=np.arange(0,len(t)) #Creater index vector
    whoops=0 #initialize error function
    for i in index: #Begin for loop which iterates over all the entries in the time vector, assigning them the value td
        if i==0: continue
        else:
            yold=yw #update the old y-value
            yw[0]=1 #hardcode in boundary condition
            out=RJ(x,yw,p); #Calculate Residual and Jacobian from new y value
            R=out[0] #Grab the Residual
            J=out[1] #Grab the Jacobian
            #print('This is the residual')
            #print(R)
            R=yw-yold-h*R
            k=0
            while np.linalg.norm(R)>tol :
                k=k+1
                J=np.eye(len(yw))-h*J; #Calculate new Jacobian from new y
                #print('This is the Jacobian and the determinant')
                #print(J)
                #print(np.linalg.det(J))
                J=sp.sparse.csc_matrix(J)
                dif=-sp.sparse.linalg.spsolve(J,R) #Apply built in sparse Linear solver to find delta from J and R
                #dif=-np.linalg.solve(J,R)  #Regular Solver. Just keeping it in there in case the sparse solver isn't working for some reason
                #print('This is delta')
                #print(dif)
                #This section is to determine if the differecne will make the working y outside the domain, and apply a different change
                #ywc=yw+dif
                
                #This is the working area
                yw=yw+dif ; #Update y
                #print('This is the new y')
                #print(yw)
                out=RJ(x,yw,p); #Calculate Residual and Jacobian from new y value
                R=out[0] #Grab the Residual
                #print('This is the residual')
                #print(R)
                J=out[1] #Grab the Jacobian
                R=yw-yold-h*R ; #Update Residual
                if k>100:
                    print('Whoops')
                    whoops=whoops+1
                    break
            y[:,i]=yw
    return (y,whoops)

"""Code Block To Run MOL for Variety of specified parameters"""

def parameter_checker(parameter_matrix,ci): #unpack paramteres and test
    
    #Calculate other internal paramters to model
    parameter_combos_count=np.shape(parameter_matrix) [0]
    c_set = [[0 for i in range(9)] for j in range(parameter_combos_count)]
    for i in np.arange(0,parameter_combos_count,1): #Begin for loop to test the different model paramters using MOL #Check if you got he upperbound right
        h=parameter_matrix[i,0] #Define timesteps to test
        tol=parameter_matrix[i,1] #Define tolerance to test
        t1=parameter_matrix[i,2] #Define initial time for this iteration
        t2=parameter_matrix[i,3] #Define final time for this iteration
        nx=int(parameter_matrix[i,4]) #Define mesh size for this iteration
        gam=parameter_matrix[i,5] #Define dimensionless ratio of diffusivities for this iteration
        beta=parameter_matrix[i,6] #Define dimensionless ratio of potentials for this iteration
        F=parameter_matrix[i,7] #Define dimensionless forward rate constant for this iteration
        Re=parameter_matrix[i,8] #Define dimensionless reverse rate constant for this iteration
        n=parameter_matrix[i,9] #Define hill coeffecient for this rest iteration
        #Calculate internal paramters to model
        x=np.linspace(0,1,nx+1) #Define x (note, if x doesn't range from 0 to 1, should edit this)
        c_set[i][8]=x #Pass along x-vector for this parameter set for plotting
        t=np.arange(t1,t2+h,h) #Define t
        c_set[i][6]=t #Pass along t-vector for this parameter set for plotting
        ny=2*nx #Solution vector size
        nt=len(t) #Define the number of timepoints
        c_set[i][7]=nt #Pass along number of time-points used for this parameter set for plotting
        y=np.zeros((ny+2,nt)) #Initialize y
        y=y+10**(-8) #Make starting values not exactly equal to zero (divide by zero erros pop up)
        p=[gam,beta,F,Re,n] #dimensionless parameter matrix
        #Run calculation for parameters of interest
        [c,whoops]=MOL(t,x,y,h,p,tol) #Find the concntration profiles in space and time using Method of Lines (MOL)
        print('you whoopsed {} many times'.format(whoops))
        
        #Unpack the data
        cb=np.zeros((nx+1,nt)) #Initalize new concentration array where bound and unbound NP concentrations are "unpacked" such that they occupy two different matrices in the same 3-D array 
        cu=np.zeros((nx+1,nt))
        xindex=np.arange(0,nx+1)
        for x_i in xindex:
            j=2*x_i #secondary index (position of unbound NP concentration in original concentration matrix)
            k=2*x_i+1 #Secondary index (position of bound NP concentration in original concentration matrix)
            cu[x_i,:]=c[j,:]
            cb[x_i,:]=c[k,:]
        c_set[i][0]=cb
        c_set[i][1]=cu
    
        #Find Average Unbound Concentration Overtime
        average_conc_overtime=np.zeros(nt)
        t2index=np.arange(0,nt)
        for t2_i in t2index:
            average_conc_overtime[t2_i]=np.average(cu[:,t2_i])
        c_set[i][2]=average_conc_overtime
        
        #Find Change in Concentration overtime
        change_in_concentration=np.zeros(nt)
        t3index=t2index[:-1].copy()
        for t3_i in t3index:
            change_in_concentration[t3_i]=(average_conc_overtime[t3_i+1]-average_conc_overtime[t3_i])/h
        c_set[i][3]=change_in_concentration
        
        #Find Average total NP concentration Overtime
        taverage_conc_overtime=np.zeros(nt)
        t4index=np.arange(0,nt)
        for t4_i in t4index:
            taverage_conc_overtime[t4_i]=np.average(cu[:,t4_i])+np.average(cb[:,t4_i])
        c_set[i][4]=taverage_conc_overtime
        
        #Find Change in Total NP Concentration overtime
        tchange_in_concentration=np.zeros(nt)
        t5index=t4index[:-1].copy()
        for t5_i in t5index:
            tchange_in_concentration[t5_i]=(taverage_conc_overtime[t5_i+1]-taverage_conc_overtime[t5_i])/h
        c_set[i][5]=tchange_in_concentration
        
            
    return c_set    
    
    
    
"""Main Code Block"""  #This code block is for entering the ranges of dimensionless numbers to test, along with specifcying the time-step size, the mesh-size, 

#Inputs Code Block
h=np.array([0.002]) #Define timesteps to test
tol=np.array([10**(-8)])  #Define the tolerance the code will run with when running Newton-Rhapson
t1=np.array([0]) #Define initialtime vector of values to test
t2=np.array([5]) #Final Time
nx=np.array([50]) #Mesh size
gam=np.array([1]) #Define dimenionless ratio of diffusivities to test
beta=np.array([1]) #Define the dimensionless ratio of potentials to test
F=np.array([1]) #Define the dimensionless forward reaction rate constant to test
Re=np.array([1]) #Define the dimensionless reverse reaction rate constant to test
n=np.array([0.85]) #Define the hill coeffecient to test
ci=10**(-8) #Define the inital concentration in the biofilm (Can't be zero, if one wants to be zero, set it to a very small number instead)



#Pack parameters to test into parameter_matrix #Might be able to get rid of index vectors
h_count=len(h) #Define number of timesteps testing
h_iv=np.arange(0,h_count,1) #Index vector for time-step
tol_count=len(tol) #Define number of timesteps testing
tol_iv=np.arange(0,tol_count,1) #Index vector for time-step
t1_count=len(t1) #Define number of inital times to test
t1_iv=np.arange(0,t1_count,1) #Index vector for initial time
t2_count=len(t2) # Define number of final times to test
t2_iv=np.arange(0,t2_count,1) #Index vector for final time
nx_count=len(nx) # Define number of mesh sizes to test
nx_iv=np.arange(0,nx_count,1) #Index vector for mesh sizes
gam_count=len(gam) # Define number of dimenionless ratio of diffusivities to test
gam_iv=np.arange(0,gam_count,1) #Index vector for dimenionless ratio of diffusivities
beta_count=len(beta) # Define number of dimensionless ratio of potentials to test
beta_iv=np.arange(0,beta_count,1) #Index vector for dimensionless ratio of potentials
F_count=len(F) # Define number of dimensionless forward reaction rate constants to test
F_iv=np.arange(0,F_count,1) #Index vector for forward reaction rate constants
Re_count=len(Re) # Define number of dimensionless reverse reaction rate constant to test
Re_iv=np.arange(0,Re_count,1) #Index vector for everse reaction rate constant
n_count=len(n) # Define number of hill coeffecient to test
n_iv=np.arange(0,n_count,1) #Index vector for hill coeffecient
parameter_combos_count=h_count*tol_count*t1_count*t2_count*nx_count*gam_count*beta_count*F_count*Re_count*n_count #Number of combination of parameters needed to test

parameter_matrix = np.zeros((parameter_combos_count,10)) #Initialize matrix with new combo of paramters in each row. 9 is because there are 9 paramteres, can change if model changes

k=0 #Begin counter for parameter matrix loop
for h_i in h:  #Loop for timestep
    for tol_i in tol: #Loop for tolerances
        for t1_i in t1:  #Loop for inital time
            for t2_i in t2:  #Loop for final time
                for nx_i in nx:  #Loop for mesh-size
                    for gam_i in gam:  #Loop for dimenionless ratio of diffusivities
                        for beta_i in beta:  #Loop for dimenionless ratio of potentials
                            for F_i in F:  #Loop for dimensionless forward reaction rate constants
                                for Re_i in Re:  #Loop for dimensionless reverse reaction rate constants
                                    for n_i in n:  #Loop for hill coeffecient
                                        parameter_matrix[k,:]=[h_i,tol_i,t1_i,t2_i,nx_i,gam_i,beta_i,F_i,Re_i,n_i]
                                        k=k+1
                                    
c_set = parameter_checker(parameter_matrix,ci) #output the set of concentration over time and space results for each set of parameters tested

    
"""Plotting (Exported to Word Document)"""

report=docx.Document()
report.add_heading('Results from most recent run of N2-0_2',0)
style=report.styles['Normal']
font=style.font
font.name='Arial'
font.size=Pt(9)    
for pc_i in np.arange(0,parameter_combos_count,1): #Begin for loop to plot the different model paramters using MOL 
    cb=c_set[pc_i][0] #Grab current bound concentration data to plot
    cu=c_set[pc_i][1] #Grab current unbound concentration data to plot
    average_conc_overtime=c_set[pc_i][2] #Grab current average concentration data to plot (Unbound NP)
    change_in_concentration=c_set[pc_i][3] #Grab current change in concentration data to plot (Unbound NP)
    taverage_conc_overtime=c_set[pc_i][4] #Grab current change in concentration data to plot (total NP)
    tchange_in_concentration=c_set[pc_i][5] #Grab current change in concentration data to plot (Unbound NP)
    npaverage_conc_overtime=np.array(taverage_conc_overtime) #Convert taverage_conc_overtime array into np.array
    logtavg_conc_overtime=np.log(npaverage_conc_overtime)  #Logarithm taverage change in concentration overtime
    nptchange_conc=np.log(tchange_in_concentration) #total change in concentration array conerted to numpy array
    logtchange_conc=np.array(nptchange_conc) #Convert average_conc_overtime array into np.array
    t=c_set[pc_i][6] #Grab time-vector for this parameter set for plotting
    nt=c_set[pc_i][7] #Grab number of time points for this parameter set for plotting
    x=c_set[pc_i][8] #Grab the position-vector for this parameter set for plotting
    break_paragraph=report.add_paragraph('___________')
    break_paragraph.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
    report.add_heading('Parameter Set %i'%pc_i,1)
    para1=report.add_paragraph(f'Step-size (h) : {parameter_matrix[pc_i,0]}     ')
    para1.add_run(f'Initial time (t1) : {parameter_matrix[pc_i,2]}     ')
    para1.add_run(f'Final time (t2) : {parameter_matrix[pc_i,3]}     ')
    para1.add_run(f'Mesh size (nx) : {parameter_matrix[pc_i,4]}')
    para2=report.add_paragraph(f'Dimensionless ratio of diffusivity (gamma) : {parameter_matrix[pc_i,5]}     ')
    para2.add_run(f'Dimensionless ratio of potential (beta): {parameter_matrix[pc_i,6]}')
    para3=report.add_paragraph(f'Dimensionless forward rate constant (F): {parameter_matrix[pc_i,7]}     ')
    para3.add_run(f'Dimensionless reverse rate constant (R): {parameter_matrix[pc_i,8]}')
    para4=report.add_paragraph(f'Hill coeffecient (n): {parameter_matrix[pc_i,9]}     ')
    para4.add_run(f'Tolerance: {parameter_matrix[pc_i,1]}')
    
    #Find relelvant maximums and minimums
    upper_1 = np.amax(cu)*1.1 #Upper bound on Unbound Concentration
    upper_2 = np.amax(average_conc_overtime)*1.1 #Upper Bound on Average Unbound Concentration
    upper_3 = np.amax(change_in_concentration)*1.1 #Upper Bound on Change in Average Concentration
    upper_4 = np.amax(cb)*1.1 #Upper bound on Bound Concentration
    upper_5 = np.amax(taverage_conc_overtime)*1.1 #Upper Bound on Average total Concentration
    upper_6 = np.amax(tchange_in_concentration)*1.1 #Upper Bound on total Change in Average Concentration
    upper_7 = np.amax(logtavg_conc_overtime) #Upper bound on log of total average NP conc
    upper_8 = np.amax(logtchange_conc) #Upper bound on log of total change in conc
    
    
    #Unbound
    #tindex_u=np.array([0,5,10,25,50,75,100,125,150,200,250]) for masnual control over timepoints plotted
    tp_u=10 #number of time points to plot
    space_u=int((nt-1)/tp_u)
    tindex_u=np.arange(0,nt,space_u)
    plt.figure(7*pc_i+0)
    for i_u in tindex_u:
        cc_u=cu[:,i_u]
        ti_u=round(t[i_u],1)
        plt.plot(x,cc_u,label='t={}'.format(ti_u))
    plt.xlim(left=0,right=1)
    plt.ylim(bottom=0,top=upper_1)
    plt.xlabel('Position',fontsize=14)
    plt.ylabel('Dimensionless Concentration',fontsize=14)
    plt.title('Dimensionless Unbound Concentration plot',fontsize=16)
    plt.xticks(fontsize=12)
    plt.yticks(fontsize=12)
    plt.legend(loc=(0.1,0.1))
    plt.savefig('Unboundplot%d.png'%pc_i)
   
    
    #Bound
    tp_b=5 #number of time points to plot
    space_b=int((nt-1)/tp_b)
    tindex_b=np.arange(0,nt,space_b)
    plt.figure(7*pc_i+1)
    for i_b in tindex_b:
        cc_b=cb[:,i_b]
        ti_b=round(t[i_b],1)
        plt.plot(x,cc_b,label='t={}'.format(ti_b))
    plt.xlim(left=0,right=1)
    plt.ylim(bottom=0,top=upper_4)
    plt.xlabel('Position',fontsize=14)
    plt.ylabel('Dimensionless Concentration',fontsize=14)
    plt.title('Dimensionless Bound Concentration plot',fontsize=16)
    plt.xticks(fontsize=12)
    plt.yticks(fontsize=12)
    plt.legend(loc=(0.1,0.1))
    plt.savefig('Boundplot%d.png'%pc_i)
    pics_paragraph1=report.add_paragraph()
    pic1=pics_paragraph1.add_run()
    pic1.add_picture('Unboundplot%d.png'%pc_i, width=docx.shared.Inches(3))
    pic2=pics_paragraph1.add_run()
    pic2.add_picture('Boundplot%d.png'%pc_i, width=docx.shared.Inches(3))
    
    #Unbound NP Average Concetration Overtime
    plt.figure(7*pc_i+2)
    plt.plot(t,average_conc_overtime)
    plt.xlim(left=parameter_matrix[pc_i,2],right=parameter_matrix[pc_i,3])
    plt.ylim(bottom=0,top=upper_2)
    plt.xlabel('Time',fontsize=14)
    plt.ylabel('Dimensionless Concentration',fontsize=14)
    plt.title('Average Dimensionless Unbound Concentration plot',fontsize=16)
    plt.xticks(fontsize=12)
    plt.yticks(fontsize=12)
    plt.savefig('averageplot%d.png'%pc_i)
    
    #Unbound NP Change in Concentration vs Concentration
    plt.figure(7*pc_i+3)
    plt.plot(average_conc_overtime,change_in_concentration)
    plt.xlim(left=0,right=upper_2)
    plt.ylim(bottom=0,top=upper_3)
    plt.xlabel('Concentration',fontsize=14)
    plt.ylabel('Dimensionless Change in Concentration',fontsize=14)
    plt.title('Unbound dC vs C plot',fontsize=16)
    plt.xticks(fontsize=12)
    plt.yticks(fontsize=12)
    plt.savefig('changeplot%d.png'%pc_i)
    pics_paragraph2=report.add_paragraph()
    pic3=pics_paragraph2.add_run()
    pic3.add_picture('averageplot%d.png'%pc_i, width=docx.shared.Inches(3))
    pic4=pics_paragraph2.add_run()
    pic4.add_picture('changeplot%d.png'%pc_i, width=docx.shared.Inches(3))
    
    #Total NP Average Concetration Overtime
    plt.figure(7*pc_i+4)
    plt.plot(t,taverage_conc_overtime)
    plt.xlim(left=parameter_matrix[pc_i,2],right=parameter_matrix[pc_i,3])
    plt.ylim(bottom=0,top=upper_5)
    plt.xlabel('Time',fontsize=14)
    plt.ylabel('Dimensionless Concentration',fontsize=14)
    plt.title('Average Dimensionless Total Concentration plot',fontsize=16)
    plt.xticks(fontsize=12)
    plt.yticks(fontsize=12)
    plt.savefig('taverageplot%d.png'%pc_i)
    
    #Total NP Change in Concentration vs Concentration
    plt.figure(7*pc_i+5)
    plt.plot(taverage_conc_overtime,tchange_in_concentration)
    plt.xlim(left=0,right=upper_5)
    plt.ylim(bottom=0,top=upper_6)
    plt.xlabel('Concentration',fontsize=14)
    plt.ylabel('Dimensionless Change in Concentration',fontsize=14)
    plt.title('Total dC vs C plot',fontsize=16)
    plt.xticks(fontsize=12)
    plt.yticks(fontsize=12)
    plt.savefig('tchangeplot%d.png'%pc_i)
    pics_paragraph3=report.add_paragraph()
    pic5=pics_paragraph3.add_run()
    pic5.add_picture('taverageplot%d.png'%pc_i, width=docx.shared.Inches(3))
    pic6=pics_paragraph3.add_run()
    pic6.add_picture('tchangeplot%d.png'%pc_i, width=docx.shared.Inches(3))
    
    #Logarithms of Total NP Change in Concentration vs Concentration
    plt.figure(7*pc_i+6)
    plt.plot(logtavg_conc_overtime,logtchange_conc)
    plt.xlim(left=0,right=upper_7)
    plt.ylim(bottom=0,top=upper_8)
    plt.xlabel('Log of Dimensionless Concentration',fontsize=14)
    plt.ylabel('Log of Dimensionless Change in Concentration',fontsize=14)
    plt.title('Total log(dC) vs log(C) plot',fontsize=16)
    plt.xticks(fontsize=12)
    plt.yticks(fontsize=12)
    plt.savefig('logtchangeplot%d.png'%pc_i)
    pics_paragraph4=report.add_paragraph()
    pic7=pics_paragraph4.add_run()
    pic7.add_picture('logtchangeplot%d.png'%pc_i, width=docx.shared.Inches(3))
    


    

#End timer
t_end=time.time()
total_time=t_end-t_start
print('Total time is {} sec'.format(total_time))

para5=report.add_paragraph(f'Time to Run : {total_time}     ')

report.save('N20_report.docx')


    

