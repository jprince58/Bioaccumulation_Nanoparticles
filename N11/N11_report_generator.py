#!/usr/bin/env python3
# -*- coding: utf-8 -*-

vn_report_generator=1.8

import os
import numpy as np
import matplotlib.pyplot as plt
import docx
from docx.shared import Pt
from datetime import datetime
import matplotlib.animation as anim
from matplotlib.animation import FuncAnimation

def plot_generator(c_set,parameter_combos_count,parameter_matrix,new_count_number,vn_N2,vn_Main_Code,vn_parameter_matrix_generator,vn_parameter_checker,vn_csv_generator,vn_method_of_lines,vn_RJ,machine_number,internal_export_path):
    """Static Plotting (Exported to Word Document)"""
    report=docx.Document()
    report.add_heading(f'Results from N6 Run #{new_count_number}-{machine_number}',0)
    now = datetime.now()
    dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
    date_time_line=report.add_paragraph('Date and Time Report Generated:  ')
    date_time_line.add_run(dt_string)
    report.add_paragraph(f'Machine number ran on- {machine_number}')
    report.add_paragraph(f'N2 version: {vn_N2}')
    report.add_paragraph(f'Main Code version: {vn_Main_Code}')
    report.add_paragraph(f'Parameter Matrix Generator version: {vn_parameter_matrix_generator}')
    report.add_paragraph(f'Parameter Checker version: {vn_parameter_checker}')
    report.add_paragraph(f'CSV Generator version: {vn_csv_generator}')
    report.add_paragraph(f'Method of Lines version: {vn_method_of_lines}')
    report.add_paragraph(f'Residual-Jacobian Calculator version: {vn_RJ}')
    report.add_paragraph(f'Report Generator version: {vn_report_generator}')
    #report.add_paragraph(f'Linear Approximator version: {vn_linear_fitting}')
    style=report.styles['Normal']
    font=style.font
    font.name='Arial'
    font.size=Pt(9)  
        
     
    for pc_i in np.arange(0,parameter_combos_count,1): #Begin for loop to plot the different model paramters using MOL 
        cm=c_set[pc_i][3] #Grab current Mobile concentration data to plot
        ca=c_set[pc_i][4] #Grab current Attached concentration data to plot
        ct=c_set[pc_i][5] #Grab current Total concentration data to plot
        cm_ss=c_set[pc_i][9]
        ca_ss=c_set[pc_i][10]
        ct_ss=c_set[pc_i][11]
        average_mobile_conc_overtime=c_set[pc_i][6]
        average_attached_conc_overtime=c_set[pc_i][7]
        average_total_conc_overtime=c_set[pc_i][8]
        t=c_set[pc_i][1] #Grab time-vector for this parameter set for plotting
        nt=c_set[pc_i][2] #Grab number of time points for this parameter set for plotting
        x=c_set[pc_i][0] #Grab the position-vector for this parameter set for plotting
        break_paragraph=report.add_paragraph('___________')
        break_paragraph.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
        """Will add these back once code is working"""
        report.add_heading('Parameter Set %i'%pc_i,1)
        para1=report.add_paragraph(f'Step-size (h) : {parameter_matrix[pc_i,0]}     ')
        para1.add_run(f'Initial time (t1) : {parameter_matrix[pc_i,2]}     ')
        para1.add_run(f'Final time (t2) : {parameter_matrix[pc_i,3]}     ')
        para1.add_run(f'Mesh size (nx) : {parameter_matrix[pc_i,4]}')
        para2=report.add_paragraph(f'Effective diffusivity (omega) : {parameter_matrix[pc_i,5]}     ')
        para2.add_run(f'Dimensionless attachment rate constant (mu): {parameter_matrix[pc_i,6]}')
        para3=report.add_paragraph(f'Dimensionless binding site density (nu): {parameter_matrix[pc_i,7]}     ')
        para3.add_run(f'Dimensionless minimum interstitial porosity (epsilon): {parameter_matrix[pc_i,8]}')
        para4=report.add_paragraph(f'Dimensionless minimum total porosity (rho): {parameter_matrix[pc_i,9]}     ')
        para4.add_run(f'Tolerance: {parameter_matrix[pc_i,1]}')
        para5=report.add_paragraph(f'Dimensionless equilibrium constant (kappa): {parameter_matrix[pc_i,10]}     ')
        para5.add_run(f'Binding Site Profile Shape Parameter (a): {parameter_matrix[pc_i,11]}')
        para5=report.add_paragraph(f'Interstitial Porosity Profile Shape Parameter (b): {parameter_matrix[pc_i,12]}     ')
        para5.add_run(f'Total Porosity Profile Shape Parameter (c): {parameter_matrix[pc_i,13]}')
        para6=report.add_paragraph(f'Partition Coeffecient (Kp): {parameter_matrix[pc_i,14]}     ')
      
        
        # %%Find relelvant maximums and minimums
        upper_1 = np.amax(cm)*1.1 #Upper bound on Mobile Concentration
        # upper_2 = np.amax(average_conc_overtime)*1.1 #Upper Bound on Average Mobile Concentration
        # upper_3 = np.amax(change_in_concentration)*1.1 #Upper Bound on Change in Average Concentration
        upper_4 = np.amax(ca)*1.1 #Upper bound on Bound Concentration
        upper_5= np.amax(ct)*1.1 #Upper bound on total concentration
        # upper_5 = np.amax(taverage_conc_overtime)*1.1 #Upper Bound on Average total Concentration
        # upper_6 = np.amax(tchange_in_concentration)*1.1 #Upper Bound on total Change in Average Concentration
        # upper_7 = np.amax(logtavg_conc_overtime)*1.1 #Upper bound on log of total average NP conc
        # upper_8 = np.amax(logtchange_conc)*1.1 #Upper bound on log of total change in conc
        # lower_7 = np.amin(logtavg_conc_overtime)*1.1 #Lower bound on log of total average NP conc
        # lower_8 = np.amin(logtavg_conc_overtime)*1.1 #Upper bound on log of total change in conc
        
        # %%Mobile
        #tindex_u=np.array([0,5,10,25,50,75,100,125,150,200,250]) for masnual control over timepoints plotted
        tp_u=10 #number of time points to plot

        #Linear discretization of plotted timepionts
        space_u=int((nt-1)/tp_u) #Linear discreitzation of timepoints
        tindex_u=np.arange(0,nt,space_u) #Linear discreitization of timepoints
        for i_u in tindex_u:
            cc_u=cm[:,i_u]
            ti_u=round(t[i_u],5)
            plt.plot(x,cc_u,label='t={}'.format(ti_u))
        """
        #Logarthmic discreitzation of plotted timepoints
        lognt_u=np.log10(nt) #Logarthmic timepoints
        logspace_u=round((lognt_u)/tp_u,10) #Logarthmic timepoints
        logtindex_u=np.arange(0,lognt_u,logspace_u) #Logarthmic timepoints
        plt.figure(9*pc_i+0)
        for logi_u in logtindex_u:
            i_u=int(10**logi_u)
            cc_u=cm[:,i_u]
            ti_u=round(t[i_u],5)
            plt.plot(x,cc_u,label='t={}'.format(ti_u))
        """
        
        plt.xlim(left=0,right=1)
        plt.ylim(bottom=0,top=upper_1)
        plt.xlabel('Position',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Dimensionless Mobile Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        plt.legend(loc=(0.1,0.1))
        Mobile_filename_partial=f'Mobileplot{pc_i}.png'
        Mobile_filename_full=os.path.join(internal_export_path,Mobile_filename_partial)
        plt.savefig(Mobile_filename_full)
        plt.close()
       
        
        # %%Attached
        tp_b=10 #number of time points to plot
        

        #Linear discretization of plotted timepionts
        space_b=int((nt-1)/tp_b) #Linear discreitzation of timepoints
        tindex_b=np.arange(0,nt,space_b) #Linear discreitzation of timepoints
        for i_b in tindex_b:
            cc_b=ca[:,i_b]
            ti_b=round(t[i_b],5)
            plt.plot(x,cc_b,label='t={}'.format(ti_b))
        """
        
        #Logarthmic discreitzation of plotted timepoints
        lognt_b=np.log10(nt) #Logarthmic timepoints
        logspace_b=round((lognt_b)/tp_b,10) #Logarthmic timepoints
        logtindex_b=np.arange(0,lognt_b,logspace_b) #Logarthmic timepoints
        plt.figure(9*pc_i+1)
        for logi_b in logtindex_b:
            i_b=int(10**logi_b)
            cc_b=ca[:,i_b]
            ti_b=round(t[i_b],5)
            plt.plot(x,cc_b,label='t={}'.format(ti_b))
 """
    
        plt.xlim(left=0,right=1)
        plt.ylim(bottom=0,top=upper_4)
        plt.xlabel('Position',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Dimensionless Attached Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        plt.legend(loc=(0.1,0.1))
        Attached_filename_partial=f'Attachedplot{pc_i}.png'
        Attached_filename_full=os.path.join(internal_export_path,Attached_filename_partial)
        plt.savefig(Attached_filename_full)
        plt.close()
        pics_paragraph1=report.add_paragraph()
        pic1=pics_paragraph1.add_run()
        pic1.add_picture(Mobile_filename_full, width=docx.shared.Inches(3))
        pic2=pics_paragraph1.add_run()
        pic2.add_picture(Attached_filename_full, width=docx.shared.Inches(3))
        
        # %% Total
        tp_t=10 #number of time points to plot
        

        #Linear discretization of plotted timepionts
        space_t=int((nt-1)/tp_t) #Linear discreitzation of timepoints
        tindex_t=np.arange(0,nt,space_t) #Linear discreitzation of timepoints
        for i_t in tindex_t:
            cc_t=ct[:,i_t]
            ti_t=round(t[i_t],5)
            plt.plot(x,cc_t,label='t={}'.format(ti_t))
        """
        
        #Logarthmic discreitzation of plotted timepoints
        lognt_t=np.log10(nt) #Logarthmic timepoints
        logspace_t=round((lognt_t)/tp_t,10) #Logarthmic timepoints
        logtindex_t=np.arange(0,lognt_t,logspace_t) #Logarthmic timepoints
        plt.figure(9*pc_i+2)
        for logi_t in logtindex_t:
            i_t=int(10**logi_t)
            cc_t=ct[:,i_t]
            ti_t=round(t[i_t],5)
            plt.plot(x,cc_t,label='t={}'.format(ti_t))
 """
    
        plt.xlim(left=0,right=1)
        plt.ylim(bottom=0,top=upper_5)
        plt.xlabel('Position',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Dimensionless Total Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        plt.legend(loc=(0.1,0.1))
        total_filename_partial=f'Totalplot{pc_i}.png'
        total_filename_full=os.path.join(internal_export_path,total_filename_partial)
        plt.savefig(total_filename_full)
        plt.close()
        
        
    
        
        # %%Mobile NP Average Concetration Overtime
        plt.figure(9*pc_i+3)
        plt.plot(t,average_mobile_conc_overtime)
        plt.xlim(left=parameter_matrix[pc_i,2],right=parameter_matrix[pc_i,3])
        #plt.xlim(left=0,right=0.0005) #Manual Override of automatic x-axis limits
        # plt.ylim(bottom=0,top=upper_2)
        plt.xlabel('Time',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Average Dimensionless Mobile Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        avgMobile_filename_partial=f'Avg Mobileplot{pc_i}.png'
        avgMobile_filename_full=os.path.join(internal_export_path,avgMobile_filename_partial)
        plt.savefig(avgMobile_filename_full)
        plt.close()
        
        pics_paragraph2=report.add_paragraph()
        pic3=pics_paragraph2.add_run()
        pic3.add_picture(total_filename_full, width=docx.shared.Inches(3))
        pic4=pics_paragraph2.add_run()
        pic4.add_picture(avgMobile_filename_full, width=docx.shared.Inches(3))
        
        # %%Attached NP Average Concetration Overtime
        plt.figure(9*pc_i+4)
        plt.plot(t,average_attached_conc_overtime)
        plt.xlim(left=parameter_matrix[pc_i,2],right=parameter_matrix[pc_i,3])
        #plt.xlim(left=0,right=0.0005) #Manual Override of automatic x-axis limits
        # plt.ylim(bottom=0,top=upper_2)
        plt.xlabel('Time',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Average Dimensionless Attached Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        avgAttached_filename_partial=f'Avg Attachedplot{pc_i}.png'
        avgAttached_filename_full=os.path.join(internal_export_path,avgAttached_filename_partial)
        plt.savefig(avgAttached_filename_full)
        plt.close()
        
        
        # %%Total NP Average Concetration Overtime
        plt.figure(9*pc_i+5)
        plt.plot(t,average_total_conc_overtime)
        plt.xlim(left=parameter_matrix[pc_i,2],right=parameter_matrix[pc_i,3])
        #plt.xlim(left=0,right=0.0005) #Manual Override of automatic x-axis limits
        # plt.ylim(bottom=0,top=upper_2)
        plt.xlabel('Time',fontsize=14)
        plt.ylabel('Dimensionless Concentration',fontsize=14)
        plt.title('Average Dimensionless Total Concentration plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        avgTotal_filename_partial=f'Avg Totalplot{pc_i}.png'
        avgTotal_filename_full=os.path.join(internal_export_path,avgTotal_filename_partial)
        plt.savefig(avgTotal_filename_full)
        plt.close()
        
        pics_paragraph3=report.add_paragraph()
        pic5=pics_paragraph3.add_run()
        pic5.add_picture(avgAttached_filename_full, width=docx.shared.Inches(3))
        pic6=pics_paragraph3.add_run()
        pic6.add_picture(avgTotal_filename_full, width=docx.shared.Inches(3))
        
        
        # # %%Steadty-state Concentration of mobile nanoparticle
        # plt.figure(9*pc_i+6)
        # plt.plot(x,cm_ss)
        # plt.xlim(left=0,right=1)
        # #plt.xlim(left=0,right=0.0005) #Manual Override of automatic x-axis limits
        # # plt.ylim(bottom=0,top=upper_2)
        # plt.xlabel('Position',fontsize=14)
        # plt.ylabel('Dimensionless Concentration',fontsize=14)
        # plt.title('Mobile SS Dimensionless Concentration plot',fontsize=16)
        # plt.xticks(fontsize=12)
        # plt.yticks(fontsize=12)
        # mobile_ss_filename_partial=f'Mobile SSplot{pc_i}.png'
        # mobile_ss_filename_full=os.path.join(internal_export_path,mobile_ss_filename_partial)
        # plt.savefig(mobile_ss_filename_full)
        # plt.close()
        
        # # %%Steadty-state Concentration of attached nanoparticle
        # plt.figure(9*pc_i+7)
        # plt.plot(x,ca_ss)
        # plt.xlim(left=0,right=1)
        # #plt.xlim(left=0,right=0.0005) #Manual Override of automatic x-axis limits
        # # plt.ylim(bottom=0,top=upper_2)
        # plt.xlabel('Position',fontsize=14)
        # plt.ylabel('Dimensionless Concentration',fontsize=14)
        # plt.title('Attached SS Dimensionless Concentration plot',fontsize=16)
        # plt.xticks(fontsize=12)
        # plt.yticks(fontsize=12)
        # attached_ss_filename_partial=f'Attached SSplot{pc_i}.png'
        # attached_ss_filename_full=os.path.join(internal_export_path,attached_ss_filename_partial)
        # plt.savefig(attached_ss_filename_full)
        # plt.close()
        
        # pics_paragraph4=report.add_paragraph()
        # pic7=pics_paragraph4.add_run()
        # pic7.add_picture(mobile_ss_filename_full, width=docx.shared.Inches(3))
        # pic8=pics_paragraph4.add_run()
        # pic8.add_picture(attached_ss_filename_full, width=docx.shared.Inches(3))
        
        # # %%Steadty-state Concentration of total nanoparticles
        # plt.figure(9*pc_i+8)
        # plt.plot(x,ct_ss)
        # plt.xlim(left=0,right=1)
        # #plt.xlim(left=0,right=0.0005) #Manual Override of automatic x-axis limits
        # # plt.ylim(bottom=0,top=upper_2)
        # plt.xlabel('Position',fontsize=14)
        # plt.ylabel('Dimensionless Concentration',fontsize=14)
        # plt.title('Total SS Dimensionless Concentration plot',fontsize=16)
        # plt.xticks(fontsize=12)
        # plt.yticks(fontsize=12)
        # total_ss_filename_partial=f'total SSplot{pc_i}.png'
        # total_ss_filename_full=os.path.join(internal_export_path,total_ss_filename_partial)
        # plt.savefig(total_ss_filename_full)
        # plt.close()
        
        # pics_paragraph5=report.add_paragraph()
        # pic9=pics_paragraph5.add_run()
        # pic9.add_picture(total_ss_filename_full, width=docx.shared.Inches(3))
        
        # %% Model fit Plot
        pics_paragraph6=report.add_paragraph()
        modelfit_filename_partial=f'Modelfitplot{pc_i}.png'
        modelfit_filename_full=os.path.join(internal_export_path,modelfit_filename_partial)
        pic6=pics_paragraph6.add_run()
        pic6.add_picture(modelfit_filename_full, width=docx.shared.Inches(6))
       
        
        """# %%Unbound NP Change in Concentration vs Concentration
        plt.figure(7*pc_i+3)
        plt.plot(average_conc_overtime,change_in_concentration)
        plt.xlim(left=0,right=upper_2)
        plt.ylim(bottom=0,top=upper_3)
        plt.xlabel('Concentration',fontsize=14)
        plt.ylabel('Dimensionless Change in Concentration',fontsize=14)
        plt.title('Unbound dC vs C plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        dCunbound_filename_partial=f'dC Unboundplot{pc_i}.png'
        dCunbound_filename_full=os.path.join(internal_export_path,dCunbound_filename_partial)
        plt.savefig(dCunbound_filename_full)
        plt.close()
        pics_paragraph2=report.add_paragraph()
        pic3=pics_paragraph2.add_run()
        pic3.add_picture(avgunbound_filename_full, width=docx.shared.Inches(3))
        pic4=pics_paragraph2.add_run()
        pic4.add_picture(dCunbound_filename_full, width=docx.shared.Inches(3))
        """
        # # %%Total NP Average Concetration Overtime
        # plt.figure(7*pc_i+4)
        # plt.plot(t,taverage_conc_overtime)
        # plt.xlim(left=parameter_matrix[pc_i,2],right=parameter_matrix[pc_i,3])
        # #plt.xlim(left=0,right=0.0005) #Manual Override of automatic x-axis limits
        # plt.ylim(bottom=0,top=upper_5)
        # plt.xlabel('Time',fontsize=14)
        # plt.ylabel('Dimensionless Concentration',fontsize=14)
        # plt.title('Average Dimensionless Total Concentration plot',fontsize=16)
        # plt.xticks(fontsize=12)
        # plt.yticks(fontsize=12)
        # avgtotalNP_filename_partial=f'Avg totolNPplot{pc_i}.png'
        # avgtotalNP_filename_full=os.path.join(internal_export_path,avgtotalNP_filename_partial)
        # plt.savefig(avgtotalNP_filename_full)
        # plt.close()
        # pic4=pics_paragraph2.add_run()
        # pic4.add_picture(avgtotalNP_filename_full, width=docx.shared.Inches(3))
        
        """# %%Total NP Change in Concentration vs Concentration
        plt.figure(7*pc_i+5)
        plt.plot(taverage_conc_overtime,tchange_in_concentration)
        plt.xlim(left=0,right=upper_5)
        plt.ylim(bottom=0,top=upper_6)
        plt.xlabel('Concentration',fontsize=14)
        plt.ylabel('Dimensionless Change in Concentration',fontsize=14)
        plt.title('Total dC vs C plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        dCtotalNP_filename_partial=f'dC totolNPplot{pc_i}.png'
        dCtotalNP_filename_full=os.path.join(internal_export_path,dCtotalNP_filename_partial)
        plt.savefig(dCtotalNP_filename_full)
        plt.close()
        pics_paragraph3=report.add_paragraph()
        pic5=pics_paragraph3.add_run()
        pic5.add_picture(avgtotalNP_filename_full, width=docx.shared.Inches(3))
        pic6=pics_paragraph3.add_run()
        pic6.add_picture(dCtotalNP_filename_full, width=docx.shared.Inches(3))
        """
        
        
        """# %%Logarithms of Total NP Change in Concentration vs Concentration
        plt.figure(7*pc_i+6)
        plt.plot(logtavg_conc_overtime,logtchange_conc)
        plt.xlim(left=lower_7,right=upper_7)
        plt.ylim(bottom=lower_8,top=upper_8) 
        plt.xlabel('Log of Dimensionless Concentration',fontsize=14)
        plt.ylabel('Log of Dimensionless Change in Concentration',fontsize=14)
        plt.title('Total log(dC) vs log(C) plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        logdCtotalNP_filename_partial=f'logdC totolNPplot{pc_i}.png'
        logdCtotalNP_filename_full=os.path.join(internal_export_path,logdCtotalNP_filename_partial)
        plt.savefig(logdCtotalNP_filename_full)
        plt.close()
        pics_paragraph4=report.add_paragraph()
        pic7=pics_paragraph4.add_run()
        pic7.add_picture(logdCtotalNP_filename_full, width=docx.shared.Inches(3))
        """
        # # %%Unbound Concentration Animation
        # unbound_anim_fig=plt.figure()
        # unbound_anim_plot=plt.plot([])
        # unbound_anim_holder=unbound_anim_plot[0]
        # plt.xlim(left=0,right=1)
        # plt.ylim(bottom=0,top=upper_1)
        # plt.xlabel('Position',fontsize=14)
        # plt.ylabel('Dimensionless Concentration',fontsize=14)
        # plt.title('Dimensionless Unbound Concentration',fontsize=16)
        # plt.xticks(fontsize=12)
        # plt.yticks(fontsize=12)
        # tp_u_anim=100 #number of time points to plot for the animation for unbound concntration
        # space_u_anim=int((nt-1)/tp_u_anim) #space between timepoints    
        # def unbound_animate(frame):
        #     #update plot
        #     c_u_plot=cu[:,frame*space_u_anim]
        #     unbound_anim_holder.set_data((x,c_u_plot))
        # unbound_anim=anim.FuncAnimation(unbound_anim_fig,unbound_animate,frames=tp_u_anim,interval=100)
        # unbound_anim_filename_partial=f'unboun_anim{pc_i}.gif'
        # unbound_anim_filename_full=os.path.join(internal_export_path,unbound_anim_filename_partial)
        # unbound_anim.save(unbound_anim_filename_full)
        # #Only need these lines if log plot is turned off
        # pics_paragraph3=report.add_paragraph()
        # pic5=pics_paragraph3.add_run()
        # #End of possibly neccesary lines
        # pic5.add_picture(unbound_anim_filename_full, width=docx.shared.Inches(3))
        # plt.close()
        
        """# %%Total NP Change in Concentration vs Concentration Animation
        dCvC_anim_fig=plt.figure()
        dCvC_anim_plot=plt.plot([])
        dCvC_anim_holder=dCvC_anim_plot[0]
        plt.xlim(left=0,right=upper_5)
        plt.ylim(bottom=0,top=upper_6)
        plt.xlabel('Concentration',fontsize=14)
        plt.ylabel('Dimensionless Change in Concentration',fontsize=14)
        plt.title('Total dC vs C plot',fontsize=16)
        plt.xticks(fontsize=12)
        plt.yticks(fontsize=12)
        tp_dCvC_anim=100 #number of time points to plot for the animation for unbound concntration
        space_dCvC_anim=int((nt-1)/tp_dCvC_anim) #space between timepoints    
        def dCvC_animate(frame):
            #update plot
            dCvC_anim_holder.set_data((taverage_conc_overtime[0:frame*space_dCvC_anim],tchange_in_concentration[0:frame*space_dCvC_anim]))
        dCvC_anim=anim.FuncAnimation(dCvC_anim_fig,dCvC_animate,frames=tp_dCvC_anim,interval=100)
        dCvC_anim_filename_partial=f'dCvC_anim{pc_i}.gif'
        dCvC_anim_filename_full=os.path.join(internal_export_path,dCvC_anim_filename_partial)
        dCvC_anim.save(dCvC_anim_filename_full)
        #pics_paragraph5=report.add_paragraph() commented out when no log plot
        #pic8=pics_paragraph5.add_run()
        pic7.add_picture(dCvC_anim_filename_full, width=docx.shared.Inches(3))
        """
        # # %%Total NP Concentration vs Time Animation
        # totCvt_anim_fig=plt.figure()
        # totCvt_anim_plot=plt.plot([])
        # totCvt_anim_holder=totCvt_anim_plot[0]
        # plt.xlim(left=parameter_matrix[pc_i,2],right=parameter_matrix[pc_i,3])
        # plt.ylim(bottom=0,top=upper_5)
        # plt.xlabel('Time',fontsize=14)
        # plt.ylabel('Dimensionless Concentration',fontsize=14)
        # plt.title('Average Dimensionless Total Concentration plot',fontsize=16)
        # plt.xticks(fontsize=12)
        # plt.yticks(fontsize=12)
        # tp_totCvt_anim=100 #number of time points to plot for the animation for unbound concntration
        # space_totCvt_anim=int((nt-1)/tp_totCvt_anim) #space between timepoints    
        # def totCvt_animate(frame):
        #     #update plot
        #     totCvt_anim_holder.set_data((t[0:frame*space_totCvt_anim],taverage_conc_overtime[0:frame*space_totCvt_anim]))
        # totCvt_anim=anim.FuncAnimation(totCvt_anim_fig,totCvt_animate,frames=tp_totCvt_anim,interval=100)
        # totCvt_anim_filename_partial=f'totCvt_anim{pc_i}.gif'
        # totCvt_anim_filename_full=os.path.join(internal_export_path,totCvt_anim_filename_partial)
        # totCvt_anim.save(totCvt_anim_filename_full)
        # #pics_paragraph5=report.add_paragraph() #Added when log plot off
        # pic6=pics_paragraph3.add_run() #Added when log plot off
        # pic6.add_picture(totCvt_anim_filename_full, width=docx.shared.Inches(3))
        # plt.close()
        
        # %% Plot Approximation for Total NP conc Overtime
        # linear_filename_partial=f'Linearplot{pc_i}.png'
        # linear_filename_full=os.path.join(internal_export_path,linear_filename_partial)
        # #pics_paragraph6=report.add_paragraph() commented out  when log plots out
        # #pic9=pics_paragraph6.add_run() commented out  when log plots out
        # pics_paragraph4=report.add_paragraph()
        # pic7=pics_paragraph4.add_run()
        # pic7.add_picture(linear_filename_full,width=docx.shared.Inches(3))
        # log_filename_partial=f'Logplot{pc_i}.png'
        # log_filename_full=os.path.join(internal_export_path,log_filename_partial)
        # pics_paragraph3=report.add_paragraph() #Added when log plot off
        # pic8=pics_paragraph3.add_run() #Added when log plot off
        # pic8.add_picture(log_filename_full,width=docx.shared.Inches(3))
        
        # # %% Add Table for Fit
        # perc_acc_table=perc_acc_matrix[pc_i][0]
        # [table_rows,table_columns]=perc_acc_table.shape
        # table1=report.add_table(rows=table_rows+1, cols=table_columns)
        # row=table1.rows[0]
        # row.cells[0].text='Percent Accumulated'
        # row.cells[1].text='Time for Model'
        # row.cells[2].text='Time for Approximation'
        # row.cells[3].text='Percent Error'
        # i_v = np.arange(0,table_rows,1) #index for rows of percent accumulation table
        # j_v = np.arange(0,table_columns,1) #inde for columns of percent accumulation table
        # for i in i_v:
        #     for j in j_v:
        #         cell=table1.cell(i+1,j)
        #         cell.text=str(perc_acc_table[i,j])
        
        # %%
    return report

# %%
"""
Created on Sat Oct 31 18:13:53 2020

@author: joshuaprince

Purpose: Script to auto-generate report from data, including plotting of key figures and generation of an animated plot

Version 1.8

Changes from Version 1.7 to 1.8 (11/13/2020 2:30 am)

Changes from Version 1.6 to 1.7 (11/11/2020 5:15 pm):
    -Plot the Approximation of file and print the table of fit
    -Added animation for dC vs C plot and Total NP Concentration Overtime

Changes from 1.5 to 1.6 (11/3/2020 5:25 pm):
    -going from linear discretization of time points to plot to logarthmic discretization

Differences between version 1.4 and 1.5 (11/3/2020 4:35 pm):
    -Added version number to script, along with reporting versions for each script

Differences between version 1.3 and 1.4 (11/3/2020 8:15 am):
    -Added plot close functionalities for all plots (python was complaining about holding so many plots in memory)
    -Lines 109, 132, 152, 167, 187, 202, 222: Added plt.close() 

Differences between 1.2 and 1.3 (11/3/2020 8:00 am)
-Changed file-path of animated unbound gif from direct export folder to internal export folder, since it is now incorporated into report
-Line 28: Got rid of direct_export_path from function call (no longer needed)
-Line 235: Changed "direct_export_path" to "internal_export_path"

Differences between 1.1 and 1.2
-Added lower bound calculations (needed for log plot), incorporated into log plotting
"""